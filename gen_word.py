import docx
from docx.shared import Inches

def create_dict_word():
    doc = docx.Document()
    doc.add_heading('Diccionario de Datos - Sistema de Gestión Documental', 0)

    tables_data = [
        ("alertas", "Indica las notificaciones para usuarios.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("uuid", "UUID", "NO", "Identificador único universal."),
            ("usuario_id", "INTEGER", "SÍ", "Destinatario de la alerta."),
            ("titulo", "VARCHAR(200)", "NO", "Título del mensaje."),
            ("mensaje", "TEXT", "NO", "Contenido de la alerta."),
            ("tipo_alerta", "VARCHAR(20)", "NO", "Categoría (info, warning)."),
            ("leida", "BOOLEAN", "NO", "Estado de lectura."),
            ("fecha_creacion", "TIMESTAMP", "NO", "Fecha de generación."),
            ("fecha_lectura", "TIMESTAMP", "SÍ", "Fecha de marcado como leída."),
            ("documento_id", "INTEGER", "SÍ", "Referencia al documento."),
            ("movimiento_id", "INTEGER", "SÍ", "Referencia al movimiento.")
        ]),
        ("anexos", "Archivos digitales adjuntos.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("documento_id", "INTEGER", "NO", "Documento al que pertenece."),
            ("nombre_archivo", "VARCHAR(255)", "NO", "Nombre del archivo subido."),
            ("extension", "VARCHAR(10)", "SÍ", "Ej: .pdf, .jpg."),
            ("tamano_bytes", "BIGINT", "SÍ", "Tamaño en bytes."),
            ("url_archivo", "VARCHAR(500)", "SÍ", "Ruta o URL de almacenamiento."),
            ("tipo_contenido", "VARCHAR(100)", "SÍ", "MIME Type."),
            ("fecha_registro", "TIMESTAMP", "NO", "Fecha de subida."),
            ("activo", "BOOLEAN", "NO", "Estado de borrado lógico.")
        ]),
        ("areas", "Departamentos de la institución.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("nombre", "VARCHAR(100)", "NO", "Nombre del área."),
            ("codigo", "VARCHAR(20)", "NO", "Sigla (ej: CONT)."),
            ("descripcion", "TEXT", "SÍ", "Detalle de funciones."),
            ("activo", "BOOLEAN", "NO", "Estado del área."),
            ("fecha_creacion", "TIMESTAMP", "NO", "Fecha de registro.")
        ]),
        ("auditoria", "Log de acciones críticas.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("usuario_id", "INTEGER", "SÍ", "Autor de la acción."),
            ("accion", "VARCHAR(100)", "NO", "Tipo: LOGIN, INSERT, etc."),
            ("tabla_afectada", "VARCHAR(50)", "SÍ", "Tabla impactada."),
            ("detalle", "TEXT", "SÍ", "Datos modificados."),
            ("ip_address", "TEXT", "SÍ", "IP del cliente."),
            ("fecha_accion", "TIMESTAMP", "NO", "Fecha y hora.")
        ]),
        ("carpetas", "Estructura de archivadores (Sprint 2).", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("nombre", "VARCHAR(100)", "NO", "Nombre del archivador."),
            ("gestion", "VARCHAR(4)", "NO", "Año fiscal (2026)."),
            ("carpeta_padre_id", "INTEGER", "SÍ", "Subcarpeta de..."),
            ("activo", "BOOLEAN", "NO", "Estado físico."),
            ("fecha_creacion", "TIMESTAMP", "NO", "Fecha de registro.")
        ]),
        ("configuracion", "Parámetros globales.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("clave", "VARCHAR(100)", "NO", "Nombre del parámetro."),
            ("valor", "TEXT", "SÍ", "Valor actual."),
            ("tipo_dato", "VARCHAR(20)", "NO", "String, Int, etc.")
        ]),
        ("documentos", "Registro principal de información.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("codigo", "VARCHAR(50)", "NO", "Código QR único."),
            ("numero_correlativo", "VARCHAR(50)", "NO", "Serie institucional."),
            ("tipo_documento_id", "INTEGER", "NO", "Referencia a tipos."),
            ("area_actual_id", "INTEGER", "NO", "Donde está ahora."),
            ("gestion", "VARCHAR(4)", "NO", "Año del trámite."),
            ("fecha_documento", "DATE", "NO", "Fecha física del papel."),
            ("estado", "VARCHAR(20)", "NO", "Activo, Archivado, etc."),
            ("nivel_confidencialidad", "INT", "NO", "Acceso del 1 al 5.")
        ]),
        ("documento_palabras_clave", "Tags de búsqueda.", [
            ("documento_id", "INTEGER", "NO", "FK Documento."),
            ("palabra_clave_id", "INTEGER", "NO", "FK Palabra Clave.")
        ]),
        ("historial_documento", "Trazabilidad de cambios.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("documento_id", "INTEGER", "NO", "Doc afectado."),
            ("fecha_cambio", "TIMESTAMP", "NO", "Cuándo ocurrió."),
            ("tipo_cambio", "VARCHAR(50)", "NO", "Acción realizada.")
        ]),
        ("movimientos", "Préstamos y derivaciones.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("documento_id", "INTEGER", "NO", "Documento en tránsito."),
            ("tipo_movimiento", "VARCHAR(20)", "NO", "Prestamo, Derivacion."),
            ("usuario_id", "INTEGER", "SÍ", "Quien lo tiene."),
            ("fecha_movimiento", "TIMESTAMP", "NO", "Salida."),
            ("fecha_devolucion", "TIMESTAMP", "SÍ", "Retorno real."),
            ("estado", "VARCHAR(20)", "NO", "Activo, Devuelto.")
        ]),
        ("palabras_clave", "Diccionario de tags.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("palabra", "VARCHAR(50)", "NO", "Etiqueta."),
            ("activo", "BOOLEAN", "NO", "Estado.")
        ]),
        ("permisos", "Acciones permitidas.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("codigo", "VARCHAR(50)", "NO", "Identificador lógico."),
            ("nombre", "VARCHAR(100)", "NO", "Descripción corta.")
        ]),
        ("rol_permisos", "Matriz de roles.", [
            ("rol", "VARCHAR(50)", "NO", "Nombre del rol."),
            ("permiso_id", "INTEGER", "NO", "FK Permiso.")
        ]),
        ("tarjetas", "ID's de personal.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("nombre_completo", "VARCHAR(255)", "NO", "Titular."),
            ("numero_identificacion", "VARCHAR(50)", "NO", "CI."),
            ("estado", "VARCHAR(20)", "NO", "Activa/Vencida.")
        ]),
        ("tipos_documento", "Maestro de tipologías.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("nombre", "VARCHAR(100)", "NO", "CI, Oficio, etc."),
            ("codigo", "VARCHAR(20)", "NO", "Abreviación.")
        ]),
        ("usuarios", "Credenciales de acceso.", [
            ("id", "SERIAL", "NO", "Clave primaria."),
            ("nombre_usuario", "VARCHAR(50)", "NO", "Login."),
            ("email", "VARCHAR(255)", "NO", "Correo."),
            ("rol", "VARCHAR(30)", "NO", "Rol asignado.")
        ]),
        ("usuario_permisos", "Excepciones de seguridad.", [
            ("usuario_id", "INTEGER", "NO", "FK Usuario."),
            ("permiso_id", "INTEGER", "NO", "FK Permiso."),
            ("denegado", "BOOLEAN", "NO", "TRUE=Quita.")
        ]),
        ("tipo_movimientos", "Catálogo flujos.", [
            ("id", "INT", "NO", "ID."),
            ("nombre", "VARCHAR(50)", "NO", "Nombre.")
        ]),
        ("notificaciones_config", "Config canales.", [
            ("id", "INT", "NO", "ID."),
            ("canal", "VARCHAR(20)", "NO", "Email/Push.")
        ])
    ]

    for table_name, desc, columns in tables_data:
        doc.add_heading(f"Tabla: {table_name}", level=1)
        doc.add_paragraph(desc)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Campo'
        hdr_cells[1].text = 'Tipo'
        hdr_cells[2].text = 'Nulo'
        hdr_cells[3].text = 'Descripción'
        
        for col, dtype, null, ddesc in columns:
            row_cells = table.add_row().cells
            row_cells[0].text = col
            row_cells[1].text = dtype
            row_cells[2].text = null
            row_cells[3].text = ddesc
        
        doc.add_paragraph()

    doc.save('DiccionarioDatos_19_Tablas.docx')
    print("Diccionario generado exitosamente en .docx")

if __name__ == "__main__":
    create_dict_word()
